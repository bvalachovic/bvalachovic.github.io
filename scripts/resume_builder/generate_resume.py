#!/usr/bin/env python3
"""
generate_resume.py — AI-powered resume generator for Brian Valachovic

Usage:
    python generate_resume.py --company "Stripe" --jd job_description.txt
    python generate_resume.py --company "Stripe"          # paste JD interactively
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

import yaml
from anthropic import Anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from dotenv import load_dotenv

# ── Paths ──────────────────────────────────────────────────────────────────────
def _find_root() -> Path:
    """Walk up from this file until we find _config.yml (the Jekyll repo root)."""
    for path in Path(__file__).resolve().parents:
        if (path / "_config.yml").exists():
            return path
    raise FileNotFoundError("Repo root not found — _config.yml not detected in any parent directory.")

ROOT = _find_root()
load_dotenv(ROOT / ".env")
RESUMES_DIR = ROOT / "_resumes"
CASE_STUDIES_DIR = ROOT / "_case_studies"
GENERATED_DIR = RESUMES_DIR / "generated"
PROFILE_PATH = RESUMES_DIR / "profile.yaml"

# ── Design tokens ──────────────────────────────────────────────────────────────
FONT = "Cambria"
COLOR_INK = RGBColor(26, 32, 44)       # near-black for name/headings
COLOR_RULE = RGBColor(44, 62, 80)      # dark slate for section rules
COLOR_MUTED = RGBColor(80, 90, 100)    # dates, secondary text


# ── Data loading ───────────────────────────────────────────────────────────────

def load_profile() -> dict:
    with open(PROFILE_PATH, encoding="utf-8") as f:
        return yaml.safe_load(f)


def load_case_studies() -> list[dict]:
    studies = []
    for path in sorted(CASE_STUDIES_DIR.glob("*.md")):
        raw = path.read_text(encoding="utf-8")
        if raw.startswith("---"):
            parts = raw.split("---", 2)
            if len(parts) >= 3:
                fm = yaml.safe_load(parts[1]) or {}
                body = parts[2].strip()
                # Skip stubs (entirely commented out)
                visible = re.sub(r"<!--.*?-->", "", body, flags=re.DOTALL).strip()
                if len(visible) < 100:
                    continue
                studies.append({
                    "title": fm.get("title", path.stem),
                    "subtitle": fm.get("subtitle", ""),
                    "role": fm.get("role", ""),
                    "content": body,
                })
        else:
            studies.append({"title": path.stem, "content": raw})
    return studies


def get_job_description(jd_path: str | None) -> str:
    if jd_path:
        return Path(jd_path).read_text(encoding="utf-8")
    print("Paste the job description below.")
    print("When done, press Enter then Ctrl+Z + Enter (Windows) or Ctrl+D (Mac/Linux):\n")
    return sys.stdin.read()


# ── Claude prompt ──────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are an expert resume writer specializing in product leadership and UX design roles.

Your task is to create a tailored, ATS-optimized resume for Brian Valachovic based on a specific job description.

Guidelines:
- Read the job description carefully to identify the role type, key responsibilities, required skills, and company context.
- Select the most relevant executive summary framing (design_leader or product_leader) — or write a custom
  blend when the role spans both disciplines. Keep it to 3–4 focused sentences.
- For each job in his experience, choose 3–5 bullets that most directly address the job requirements.
  Prioritize bullets with quantified outcomes ($, %, time saved). Do not include bullets that are
  clearly irrelevant to the target role.
- Pull in specific metrics or outcomes from the case studies where they add meaningful context beyond
  the work experience bullets. Do not duplicate content already covered in a bullet.
- Do NOT fabricate or exaggerate any figures — only use numbers already present in the profile or case studies.
- Emphasize skills and tools that appear in or closely relate to the job description.
- Adjust key_differentiators to surface the 4–5 most relevant ones for this specific role.
- Keep each bullet to 1–2 lines. Use strong action verbs (past tense for past roles, present for current).
- The role_title field is the tag line under Brian's name — tailor it to the target role (e.g. "Head of Product Design  |  Principal UX Designer").

Return ONLY valid JSON — no markdown fences, no commentary — matching this exact schema:

{
  "role_title": "string — tailored title line for this application",
  "summary": "string — 3–4 sentence executive summary tailored to this role",
  "experience": [
    {
      "company": "string — company name and location",
      "title": "string — job title",
      "dates": "string — date range",
      "description": "string — one-line role context (optional, keep brief or omit if redundant)",
      "bullets": ["string", "string", "..."]
    }
  ],
  "earlier_roles": "string — single line summarizing earlier roles",
  "skills": {
    "product_management": "string — comma-separated, most relevant first",
    "ux_design": "string — comma-separated, most relevant first",
    "technical": "string — comma-separated, most relevant first",
    "tools": "string — comma-separated, most relevant first"
  },
  "education": ["string", "string"],
  "key_differentiators": ["string — full label: description line", "..."]
}"""


def build_user_message(profile: dict, case_studies: list[dict], company: str, jd: str) -> str:
    profile_text = yaml.dump(profile, default_flow_style=False, allow_unicode=True)
    studies_text = "\n\n---\n\n".join(
        f"### {cs['title']}\n**Role context:** {cs.get('role', 'N/A')}\n\n{cs['content']}"
        for cs in case_studies
    )
    return f"""Here is Brian's complete professional profile:

```yaml
{profile_text}
```

Here are his portfolio case studies with detailed outcomes:

{studies_text}

---

Target company: {company}

Job description:
{jd}

Please generate the tailored resume JSON now."""


def call_claude(profile: dict, case_studies: list[dict], company: str, jd: str) -> dict:
    client = Anthropic()
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": build_user_message(profile, case_studies, company, jd)}],
    )
    text = response.content[0].text.strip()

    # Strip markdown code fences if Claude included them despite instructions
    if text.startswith("```"):
        lines = text.splitlines()
        text = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])

    try:
        return json.loads(text)
    except json.JSONDecodeError as e:
        print(f"\nError: Claude returned invalid JSON.\nRaw response:\n{text}")
        raise SystemExit(1) from e


# ── Word document generation ───────────────────────────────────────────────────

def _set_font(run, size: float, bold=False, italic=False, color: RGBColor | None = None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _bottom_rule(paragraph, color: str = "2C3E50", width: str = "6"):
    """Add a bottom border line to a paragraph (used as section divider)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), width)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_header(doc: Document, label: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(label.upper())
    _set_font(run, 10.5, bold=True, color=COLOR_RULE)
    _bottom_rule(p)
    return p


def _bullet_para(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    _set_font(run, 10)
    return p


def build_docx(data: dict, output_path: Path):
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    # ── Header block ──────────────────────────────────────────────────────────
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after = Pt(2)
    _set_font(name_p.add_run("Brian Valachovic"), 24, bold=True, color=COLOR_INK)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    _set_font(title_p.add_run(data["role_title"]), 11.5, color=COLOR_MUTED)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after = Pt(0)
    _set_font(contact_p.add_run("bvalachovic@outlook.com  |  609.864.4248"), 10, color=COLOR_MUTED)

    # ── Executive Summary ─────────────────────────────────────────────────────
    _section_header(doc, "Executive Summary")
    summary_p = doc.add_paragraph()
    summary_p.paragraph_format.space_after = Pt(2)
    _set_font(summary_p.add_run(data["summary"]), 10)

    # ── Experience ────────────────────────────────────────────────────────────
    _section_header(doc, "Experience")
    for job in data["experience"]:
        company_p = doc.add_paragraph()
        company_p.paragraph_format.space_before = Pt(7)
        company_p.paragraph_format.space_after = Pt(0)
        _set_font(company_p.add_run(job["company"]), 10.5, bold=True)

        td_p = doc.add_paragraph()
        td_p.paragraph_format.space_before = Pt(0)
        td_p.paragraph_format.space_after = Pt(2)
        _set_font(td_p.add_run(job["title"]), 10, bold=True, italic=True)
        _set_font(td_p.add_run(f"  |  {job['dates']}"), 10, color=COLOR_MUTED)

        # Optional one-line role description
        desc = job.get("description", "").strip()
        if desc:
            desc_p = doc.add_paragraph()
            desc_p.paragraph_format.space_before = Pt(0)
            desc_p.paragraph_format.space_after = Pt(2)
            _set_font(desc_p.add_run(desc), 10, italic=True, color=COLOR_MUTED)

        for bullet in job["bullets"]:
            _bullet_para(doc, bullet)

    # Earlier roles
    earlier_p = doc.add_paragraph()
    earlier_p.paragraph_format.space_before = Pt(7)
    earlier_p.paragraph_format.space_after = Pt(0)
    _set_font(earlier_p.add_run("Earlier Roles:  "), 10, bold=True)
    _set_font(earlier_p.add_run(data["earlier_roles"]), 10)

    # ── Skills ────────────────────────────────────────────────────────────────
    _section_header(doc, "Technical & Product Expertise")
    skill_rows = [
        ("Product Management", data["skills"].get("product_management", "")),
        ("UX & Design", data["skills"].get("ux_design", "")),
        ("Technical", data["skills"].get("technical", "")),
        ("Tools", data["skills"].get("tools", "")),
    ]
    for label, content in skill_rows:
        if not content:
            continue
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(1)
        sp.paragraph_format.space_after = Pt(1)
        _set_font(sp.add_run(f"{label}: "), 10, bold=True)
        _set_font(sp.add_run(content), 10)

    # ── Education ─────────────────────────────────────────────────────────────
    _section_header(doc, "Education & Certifications")
    edu_p = doc.add_paragraph()
    edu_p.paragraph_format.space_after = Pt(2)
    _set_font(edu_p.add_run("  •  ".join(data["education"])), 10)

    # ── Key Differentiators ───────────────────────────────────────────────────
    _section_header(doc, "Key Differentiators")
    for diff in data["key_differentiators"]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if ":" in diff:
            label, rest = diff.split(":", 1)
            _set_font(p.add_run(label + ":"), 10, bold=True)
            _set_font(p.add_run(rest), 10)
        else:
            _set_font(p.add_run(diff), 10)

    doc.save(output_path)
    print(f"\nResume saved to:  {output_path.relative_to(ROOT)}")


# ── Entry point ────────────────────────────────────────────────────────────────

def slugify(name: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-")


def main():
    parser = argparse.ArgumentParser(description="Generate a tailored resume for a job application.")
    parser.add_argument("--company", required=True, help='Company name, e.g. "Stripe"')
    parser.add_argument("--jd", help="Path to a .txt file containing the job description (omit to paste interactively)")
    args = parser.parse_args()

    if not os.getenv("ANTHROPIC_API_KEY"):
        print("Error: ANTHROPIC_API_KEY not set. Copy .env.example to .env and add your key.")
        sys.exit(1)

    GENERATED_DIR.mkdir(parents=True, exist_ok=True)

    print("Loading profile...")
    profile = load_profile()

    print("Loading case studies...")
    case_studies = load_case_studies()
    print(f"  {len(case_studies)} case studies loaded.")

    jd = get_job_description(args.jd)
    if not jd.strip():
        print("Error: job description is empty.")
        sys.exit(1)

    filename = f"brian-valachovic-resume-{slugify(args.company)}.docx"
    output_path = GENERATED_DIR / filename

    print(f"\nCalling Claude to tailor resume for {args.company}...")
    data = call_claude(profile, case_studies, args.company, jd)

    print("Building Word document...")
    build_docx(data, output_path)


if __name__ == "__main__":
    main()
